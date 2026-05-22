import sys
import json
import re
from docx import Document

# 确保 Windows 控制台 UTF-8 输出
if sys.platform == 'win32':
    sys.stdout.reconfigure(encoding='utf-8')  # type: ignore[attr-defined]
    sys.stderr.reconfigure(encoding='utf-8')  # type: ignore[attr-defined]

def extract_raw_content(doc):
    """提取文档中所有原始内容"""
    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        style_name = para.style.name if para.style else ""
        if text:
            paragraphs.append({
                "text": text,
                "style": style_name,
                "is_heading": style_name.startswith("Heading") if style_name else False,
            })
    return paragraphs

def extract_tables_raw(doc):
    """提取表格内容（有些试卷把题目放表格中）"""
    tables = []
    for i, table in enumerate(doc.tables):
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(cells)
        tables.append(rows)
    return tables

def estimate_parse_quality(doc):
    """估计解析质量"""
    total_text = "\n".join(p.text for p in doc.paragraphs)
    char_count = len(total_text)
    question_marks = total_text.count("？") + total_text.count("?")
    number_patterns = len(re.findall(r'[一二三四五六七八九十]+[、．.]', total_text))

    quality = "high"
    notes = []

    if char_count < 200:
        quality = "low"
        notes.append("文档内容过少，可能为扫描件或格式异常")
    elif char_count < 500:
        quality = "medium"
        notes.append("文档内容较少，可能部分信息缺失")

    if question_marks < 2 and number_patterns < 2:
        if quality == "high":
            quality = "medium"
        notes.append("未检测到明显题目特征，可能格式不规范")

    return quality, "; ".join(notes) if notes else "解析质量良好"

def parse_docx(filepath):
    """主解析函数"""
    doc = Document(filepath)
    paragraphs = extract_raw_content(doc)
    tables = extract_tables_raw(doc)
    quality, note = estimate_parse_quality(doc)

    # 提取文档标题（第一段或第一个 Heading）
    title = ""
    for p in paragraphs:
        if p["is_heading"] or p["style"] in ("Title", "title"):
            title = p["text"]
            break
    if not title and paragraphs:
        title = paragraphs[0]["text"]

    result = {
        "title": title,
        "parse_quality": quality,
        "parse_quality_note": note,
        "paragraphs": paragraphs,
        "tables": tables,
    }

    return result

def main():
    if len(sys.argv) < 3:
        print("Usage: python parse_docx.py <input.docx> <output_raw.json>")
        sys.exit(1)

    input_path = sys.argv[1]
    output_path = sys.argv[2]

    result = parse_docx(input_path)

    with open(output_path, "w", encoding="utf-8") as f:
        json.dump(result, f, ensure_ascii=False, indent=2)

    print(f"✅ 解析完成: {len(result['paragraphs'])} 个段落 → {output_path}")
    print(f"   解析质量: {result['parse_quality']} - {result['parse_quality_note']}")
    print(f"   提示: 请交由 Agent 进行结构化处理 → clean.json")

if __name__ == "__main__":
    main()
